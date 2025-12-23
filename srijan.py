import io
from typing import List, Dict, Any

import streamlit as st

# Optional dependencies (for export)
try:
    from docx import Document  # python-docx
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except Exception:
    canvas = None


# ----------------------- Helper Functions -----------------------

def build_structured_questions(
    exam_name: str,
    class_name: str,
    subject: str,
    subtopics: List[str],
    typologies: List[str],
    difficulty: str,
    need_diagram: bool,
    num_questions: int,
    ocr_text: str,
    base_prompt: str,
    engine_config: Dict[str, Any],
) -> List[Dict[str, Any]]:
    """Stub: replace this with real dual‑engine AI calls.

    For now, returns dummy questions so the UI is usable immediately.
    Later you can plug in OpenAI / Gemini / Claude using engine_config.
    """
    questions = []
    for i in range(1, num_questions + 1):
        q = {
            "id": i,
            "question_text": (
                f"[{exam_name} | {subject}] Q{i}: Sample question on "
                f"{', '.join(subtopics) or 'chapter core concepts'} "
                f"({difficulty}, {', '.join(typologies) or 'Mixed type'})."
            ),
            "needs_diagram": need_diagram,
            "solution": "Sample solution steps will appear here once AI is wired.",
        }
        questions.append(q)
    return questions


def questions_to_plain_text(questions: List[Dict[str, Any]]) -> str:
    lines = []
    for q in questions:
        lines.append(f"Q{q['id']}. {q['question_text']}")
        if q.get("needs_diagram"):
            lines.append("[Diagram required]")
        lines.append("")
        lines.append("Solution:")
        lines.append(q.get("solution", ""))
        lines.append("\n" + "-" * 80 + "\n")
    return "\n".join(lines)


def export_to_docx(questions: List[Dict[str, Any]]) -> io.BytesIO:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Please add it to requirements.txt.")

    doc = Document()
    doc.add_heading("Generated Questions", level=1)

    for q in questions:
        doc.add_paragraph(f"Q{q['id']}. {q['question_text']}")
        if q.get("needs_diagram"):
            doc.add_paragraph("[Diagram required]", style="Intense Quote")
        doc.add_paragraph("Solution:", style="Heading 3")
        doc.add_paragraph(q.get("solution", ""))
        doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_to_pdf(questions: List[Dict[str, Any]]) -> io.BytesIO:
    if canvas is None:
        raise RuntimeError("reportlab is not installed. Please add it to requirements.txt.")

    buffer = io.BytesIO()
    c = canvas(buffer, pagesize=letter)
    width, height = letter

    text_obj = c.beginText()
    text_obj.setTextOrigin(40, height - 40)
    text_obj.setLeading(16)

    text_obj.textLine("Generated Questions")
    text_obj.textLine("" )

    plain = questions_to_plain_text(questions)
    for line in plain.splitlines():
        if text_obj.getY() <= 40:
            c.drawText(text_obj)
            c.showPage()
            text_obj = c.beginText(40, height - 40)
            text_obj.setLeading(16)
        text_obj.textLine(line)

    c.drawText(text_obj)
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer


# ----------------------- Streamlit UI -----------------------

st.set_page_config(
    page_title="New Question Generator",
    page_icon="🧠",
    layout="wide",
)

st.title("🧠 New Question Generator")
st.caption(
    "Design questions aligned with PYQ + practice sets using dual AI engines."
)

# Keep questions in session so export buttons work after generation
if "generated_questions" not in st.session_state:
    st.session_state.generated_questions = []


# --------- Layout: Inputs (left) | Engine + Output (right) ---------

col_left, col_right = st.columns([2, 1], gap="large")

with col_left:
    st.subheader("1️⃣ Exam & Question Blueprint")

    with st.form("question_spec_form"):
        exam_name = st.text_input("Exam name", placeholder="JEE Main, NEET UG, SSC CGL, ...")
        class_name = st.text_input("Class / Grade", placeholder="11, 12, Droppers, ...")
        subject = st.text_input("Subject", placeholder="Physics, Chemistry, Biology, ...")

        subtopics_raw = st.text_input(
            "Subtopics (comma‑separated)",
            placeholder="Kinematics, NLM, Work-Energy, ...",
        )
        subtopics = [s.strip() for s in subtopics_raw.split(",") if s.strip()]

        typologies_options = [
            "Single correct MCQ",
            "Multiple correct MCQ",
            "Integer type",
            "Assertion-Reason",
            "Match the columns",
            "Numerical",
            "Subjective (short)",
            "Subjective (long)",
        ]
        typologies = st.multiselect(
            "Question typologies",
            typologies_options,
            default=["Single correct MCQ"],
        )

        difficulty = st.select_slider(
            "Difficulty level",
            options=["Easy", "Easy-Medium", "Medium", "Medium-Hard", "Hard"],
            value="Medium",
        )

        need_diagram = st.radio(
            "Need diagram?",
            ["No", "Yes"],
            index=0,
            horizontal=True,
        ) == "Yes"

        num_questions = st.number_input(
            "How many questions to generate?",
            min_value=1,
            max_value=50,
            value=5,
            step=1,
        )

        st.markdown("---")
        st.subheader("2️⃣ Source Documents (OCR Output)")
        st.caption("Upload the *OCR-done* text from PYQ + practice sets (up to 2 files).")

        ocr_files = st.file_uploader(
            "Upload OCR text files",
            type=["txt"],
            accept_multiple_files=True,
            help="For now, use .txt exports from Google Vision / OCR.",
        )

        base_prompt = st.text_area(
            "Generation prompt / template",
            height=180,
            placeholder=(
                "Paste the master prompt that instructs the AI how to create totally new,\n"
                "pattern-aligned questions + solutions (teacher-grade explanation, etc.)."
            ),
        )

        submitted = st.form_submit_button("🚀 Generate Questions")

        if submitted:
            if not exam_name or not subject:
                st.error("Please fill at least *Exam name* and *Subject*.")
            elif not ocr_files:
                st.error("Please upload at least one OCR text file.")
            elif not base_prompt.strip():
                st.error("Please paste your generation prompt / template.")
            else:
                # Concatenate OCR text
                all_ocr_text = "\n\n".join(
                    f.read().decode("utf-8", errors="ignore") for f in ocr_files
                )

                engine_config = st.session_state.get("engine_config", {})

                with st.spinner("Calling AI engines and building questions..."):
                    questions = build_structured_questions(
                        exam_name=exam_name,
                        class_name=class_name,
                        subject=subject,
                        subtopics=subtopics,
                        typologies=typologies,
                        difficulty=difficulty,
                        need_diagram=need_diagram,
                        num_questions=num_questions,
                        ocr_text=all_ocr_text,
                        base_prompt=base_prompt,
                        engine_config=engine_config,
                    )

                st.session_state.generated_questions = questions
                st.success(f"Generated {len(questions)} question(s).")


with col_right:
    st.subheader("3️⃣ AI Engine Settings")

    if "engine_config" not in st.session_state:
        st.session_state.engine_config = {
            "primary_provider": "OpenAI",
            "primary_model": "gpt-4.1",
            "secondary_provider": "None",
            "secondary_model": "",
            "api_keys": {},
        }

    engine_config = st.session_state.engine_config

    primary_provider = st.selectbox(
        "Primary engine",
        ["OpenAI", "Gemini", "Claude"],
        index=["OpenAI", "Gemini", "Claude"].index(engine_config["primary_provider"]),
    )

    if primary_provider == "OpenAI":
        primary_model = st.selectbox(
            "Primary model",
            ["gpt-4.1", "gpt-4.1-mini", "gpt-4o", "gpt-5.1"],
            index=0,
        )
    elif primary_provider == "Gemini":
        primary_model = st.selectbox(
            "Primary model",
            ["gemini-1.5-pro", "gemini-1.5-flash"],
            index=0,
        )
    else:
        primary_model = st.selectbox(
            "Primary model",
            ["claude-3-5-sonnet", "claude-3-opus"],
            index=0,
        )

    primary_key = st.text_input(
        f"{primary_provider} API key",
        type="password",
        help="Key is stored only in session_state in this app.",
    )

    st.markdown("---")
    st.caption("Optional: configure a *secondary* engine for dual-model workflows.")

    use_secondary = st.checkbox("Enable dual engine (secondary model)")

    secondary_provider = "None"
    secondary_model = ""
    secondary_key = ""

    if use_secondary:
        secondary_provider = st.selectbox(
            "Secondary engine",
            ["OpenAI", "Gemini", "Claude"],
            index=1,
            key="secondary_provider_box",
        )

        if secondary_provider == "OpenAI":
            secondary_model = st.selectbox(
                "Secondary model",
                ["gpt-4.1", "gpt-4.1-mini", "gpt-4o", "gpt-5.1"],
                index=1,
                key="secondary_model_box",
            )
        elif secondary_provider == "Gemini":
            secondary_model = st.selectbox(
                "Secondary model",
                ["gemini-1.5-pro", "gemini-1.5-flash"],
                index=1,
                key="secondary_model_box",
            )
        else:
            secondary_model = st.selectbox(
                "Secondary model",
                ["claude-3-5-sonnet", "claude-3-opus"],
                index=0,
                key="secondary_model_box",
            )

        secondary_key = st.text_input(
            f"{secondary_provider} API key",
            type="password",
            key="secondary_key_box",
        )

    # Save engine config
    engine_config.update(
        {
            "primary_provider": primary_provider,
            "primary_model": primary_model,
            "secondary_provider": secondary_provider,
            "secondary_model": secondary_model,
            "api_keys": {
                primary_provider: primary_key,
                secondary_provider: secondary_key,
            },
        }
    )
    st.session_state.engine_config = engine_config

    st.markdown("---")
    st.subheader("4️⃣ Export Questions")

    questions = st.session_state.get("generated_questions", [])

    if not questions:
        st.info("Generate some questions first to enable export.")
    else:
        txt_data = questions_to_plain_text(questions)

        st.download_button(
            "📄 Download as TXT",
            data=txt_data,
            file_name="questions.txt",
            mime="text/plain",
        )

        # DOCX export
        try:
            docx_buffer = export_to_docx(questions)
            st.download_button(
                "📝 Download as DOCX",
                data=docx_buffer,
                file_name="questions.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            st.warning(f"DOCX export unavailable: {e}")

        # PDF export
        try:
            pdf_buffer = export_to_pdf(questions)
            st.download_button(
                "📕 Download as PDF",
                data=pdf_buffer,
                file_name="questions.pdf",
                mime="application/pdf",
            )
        except Exception as e:
            st.warning(f"PDF export unavailable: {e}")


st.markdown("---")
st.caption("Made for internal question-generation pipelines · Plug in your own AI calls inside `build_structured_questions()`.\n")